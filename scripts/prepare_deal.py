#!/usr/bin/env python3
"""
prepare_deal.py — Pre-processing script for Markup.

Takes a loan agreement (.docx or .txt) and creates the provision-based
folder structure that Claude Code will operate on.

Usage:
    python scripts/prepare_deal.py agreement.docx --posture borrower_friendly
    python scripts/prepare_deal.py agreement.docx --pattern "ARTICLE" --posture balanced
    python scripts/prepare_deal.py agreement.txt --posture borrower_friendly

Features:
    - Splits agreement into provisions (by heading style or text pattern)
    - Creates structured folders with manifest files
    - Extracts full agreement text for context
    - Generates review_config.json
    - Supports resume (won't overwrite existing reviewed provisions)
"""

import argparse
import json
import os
import re
import sys
import hashlib
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional


# ---------------------------------------------------------------------------
# Text-based splitting (works with .txt files)
# ---------------------------------------------------------------------------

# Common article/section patterns in loan agreements
DEFAULT_PATTERNS = [
    r'^ARTICLE\s+[IVXLCDM]+',          # ARTICLE I, ARTICLE II, etc.
    r'^ARTICLE\s+\d+',                   # ARTICLE 1, ARTICLE 2, etc.
    r'^Article\s+[IVXLCDM]+',           # Article I, Article II
    r'^Article\s+\d+',                   # Article 1, Article 2
    r'^SECTION\s+\d+[\.:]\s',           # SECTION 1. / SECTION 1:
    r'^\d+\.\s+[A-Z][A-Z\s]+$',        # "1. DEFINITIONS" style
]


def detect_split_pattern(text: str) -> Optional[str]:
    """Auto-detect the most likely section pattern in the agreement."""
    for pattern in DEFAULT_PATTERNS:
        matches = re.findall(pattern, text, re.MULTILINE)
        if len(matches) >= 3:  # Need at least 3 matches to be confident
            return pattern
    return None


def split_text_by_pattern(text: str, pattern: str) -> list[dict]:
    """
    Split agreement text into provisions based on a regex pattern.
    Returns list of dicts with 'title', 'text', 'start_line'.
    """
    lines = text.split('\n')
    regex = re.compile(pattern, re.IGNORECASE)

    split_points = []
    for i, line in enumerate(lines):
        stripped = line.strip()
        if regex.match(stripped):
            split_points.append((i, stripped))

    if not split_points:
        return []

    provisions = []

    # Preamble (content before first split point)
    if split_points[0][0] > 0:
        preamble_lines = lines[:split_points[0][0]]
        preamble_text = '\n'.join(preamble_lines).strip()
        if preamble_text:
            provisions.append({
                'title': 'Preamble',
                'text': preamble_text,
                'start_line': 0,
                'section_number': '00',
            })

    # Each provision
    for idx, (start_line, title) in enumerate(split_points):
        if idx + 1 < len(split_points):
            end_line = split_points[idx + 1][0]
        else:
            end_line = len(lines)

        provision_lines = lines[start_line:end_line]
        provision_text = '\n'.join(provision_lines).strip()

        # Generate a clean section number
        section_num = f"{idx + 1:02d}"

        provisions.append({
            'title': title.strip(),
            'text': provision_text,
            'start_line': start_line,
            'section_number': section_num,
        })

    return provisions


# ---------------------------------------------------------------------------
# DOCX-based splitting
# ---------------------------------------------------------------------------

def split_docx_by_style(docx_path: str, style_name: str) -> list[dict]:
    """Split a .docx file by heading style. Returns text-based provisions."""
    try:
        from docx import Document
    except ImportError:
        print("Error: python-docx required. Install with:")
        print("  pip install python-docx --break-system-packages")
        sys.exit(1)

    doc = Document(docx_path)
    split_points = []

    for i, para in enumerate(doc.paragraphs):
        if para.style and para.style.name.lower() == style_name.lower():
            split_points.append((i, para.text.strip()))

    if not split_points:
        return []

    provisions = []
    paragraphs = doc.paragraphs

    # Preamble
    if split_points[0][0] > 0:
        preamble_paras = paragraphs[:split_points[0][0]]
        preamble_text = '\n\n'.join(p.text for p in preamble_paras if p.text.strip())
        if preamble_text:
            provisions.append({
                'title': 'Preamble',
                'text': preamble_text,
                'start_line': 0,
                'section_number': '00',
            })

    for idx, (start_para, title) in enumerate(split_points):
        if idx + 1 < len(split_points):
            end_para = split_points[idx + 1][0]
        else:
            end_para = len(paragraphs)

        section_paras = paragraphs[start_para:end_para]
        section_text = '\n\n'.join(p.text for p in section_paras if p.text.strip())

        section_num = f"{idx + 1:02d}"
        provisions.append({
            'title': title,
            'text': section_text,
            'start_line': start_para,
            'section_number': section_num,
        })

    return provisions


def split_docx_by_pattern(docx_path: str, pattern: str) -> list[dict]:
    """Split a .docx by text pattern matching. Returns text-based provisions."""
    try:
        from docx import Document
    except ImportError:
        print("Error: python-docx required. Install with:")
        print("  pip install python-docx --break-system-packages")
        sys.exit(1)

    doc = Document(docx_path)
    regex = re.compile(pattern, re.IGNORECASE)
    split_points = []

    for i, para in enumerate(doc.paragraphs):
        if regex.match(para.text.strip()):
            split_points.append((i, para.text.strip()))

    if not split_points:
        return []

    provisions = []
    paragraphs = doc.paragraphs

    # Preamble
    if split_points[0][0] > 0:
        preamble_paras = paragraphs[:split_points[0][0]]
        preamble_text = '\n\n'.join(p.text for p in preamble_paras if p.text.strip())
        if preamble_text:
            provisions.append({
                'title': 'Preamble',
                'text': preamble_text,
                'start_line': 0,
                'section_number': '00',
            })

    for idx, (start_para, title) in enumerate(split_points):
        if idx + 1 < len(split_points):
            end_para = split_points[idx + 1][0]
        else:
            end_para = len(paragraphs)

        section_paras = paragraphs[start_para:end_para]
        section_text = '\n\n'.join(p.text for p in section_paras if p.text.strip())

        section_num = f"{idx + 1:02d}"
        provisions.append({
            'title': title,
            'text': section_text,
            'start_line': start_para,
            'section_number': section_num,
        })

    return provisions


def extract_full_text_from_docx(docx_path: str) -> str:
    """Extract all text from a .docx for the full_agreement.txt context file."""
    try:
        from docx import Document
    except ImportError:
        print("Error: python-docx required.")
        sys.exit(1)

    doc = Document(docx_path)
    return '\n\n'.join(p.text for p in doc.paragraphs if p.text.strip())


# ---------------------------------------------------------------------------
# Folder structure creation
# ---------------------------------------------------------------------------

def sanitize_folder_name(title: str) -> str:
    """Convert a provision title to a safe folder name."""
    name = title.lower()
    name = re.sub(r'[^a-z0-9\s_-]', '', name)
    name = re.sub(r'\s+', '_', name.strip())
    name = re.sub(r'_+', '_', name)
    return name[:80]


def detect_cross_references(text: str) -> list[str]:
    """Extract cross-references from provision text."""
    patterns = [
        r'Section\s+\d+[\.\d]*',
        r'Article\s+[IVXLCDM]+',
        r'Article\s+\d+',
        r'Exhibit\s+[A-Z]',
        r'Schedule\s+\d+[\.\d]*',
        r'Annex\s+[A-Z]',
    ]
    refs = set()
    for pattern in patterns:
        for match in re.finditer(pattern, text, re.IGNORECASE):
            refs.add(match.group())
    return sorted(refs)


def detect_defined_terms(text: str) -> list[str]:
    """Extract likely defined terms (capitalized multi-word phrases in quotes or caps)."""
    patterns = [
        r'"([A-Z][A-Za-z\s]+)"',         # "Defined Term"
        r'\u201c([A-Z][A-Za-z\s]+)\u201d',  # "Defined Term" (smart quotes)
    ]
    terms = set()
    for pattern in patterns:
        for match in re.finditer(pattern, text):
            term = match.group(1).strip()
            if len(term) > 2 and len(term) < 60:
                terms.add(term)
    return sorted(terms)


def create_provision_folder(
    output_dir: Path,
    provision: dict,
    full_agreement_hash: str,
) -> Path:
    """Create a provision folder with original.txt and manifest.json."""
    folder_name = f"{provision['section_number']}_{sanitize_folder_name(provision['title'])}"
    folder_path = output_dir / "provisions" / folder_name
    folder_path.mkdir(parents=True, exist_ok=True)

    # Write original text
    original_path = folder_path / "original.txt"
    original_path.write_text(provision['text'], encoding='utf-8')

    # Detect cross-references and defined terms
    cross_refs = detect_cross_references(provision['text'])
    defined_terms = detect_defined_terms(provision['text'])

    # Create manifest
    manifest = {
        "section_number": provision['section_number'],
        "title": provision['title'],
        "start_line": provision['start_line'],
        "char_count": len(provision['text']),
        "word_count": len(provision['text'].split()),
        "cross_references": cross_refs,
        "defined_terms_referenced": defined_terms,
        "status": "pending",
        "created_at": datetime.now(timezone.utc).isoformat(),
        "reviewed_at": None,
        "agreement_hash": full_agreement_hash,
        "cross_ref_flags": [],
        "open_issues": [],
    }

    manifest_path = folder_path / "manifest.json"
    manifest_path.write_text(json.dumps(manifest, indent=2), encoding='utf-8')

    return folder_path


def create_review_config(output_dir: Path, posture: str, notes: str = "",
                        has_term_sheet: bool = False,
                        skills: list[dict] = None) -> Path:
    """Create the review configuration file."""
    config = {
        "review_posture": posture,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "client_notes": notes,
        "has_term_sheet": has_term_sheet,
        "skills": [{"name": s["name"], "path": s["path"], "description": s["description"]}
                   for s in (skills or [])],
        "preferences": {
            "cure_period_minimum_days": 30 if posture == "borrower_friendly" else 10,
            "notice_period_minimum_days": 10 if posture == "borrower_friendly" else 5,
            "flag_springing_recourse": True,
            "flag_cash_management_triggers": posture == "borrower_friendly",
            "flag_transfer_restrictions": posture == "borrower_friendly",
            "require_lender_reasonableness": posture == "borrower_friendly",
            "moderation_level": "aggressive" if posture in ["borrower_friendly", "lender_friendly"] else "moderate",
        },
        "output_format": {
            "include_analysis": True,
            "include_revised_text": True,
            "include_changes_summary": True,
            "include_redline_markup": True,
        },
    }

    config_path = output_dir / "review_config.json"
    config_path.write_text(json.dumps(config, indent=2), encoding='utf-8')
    return config_path


# ---------------------------------------------------------------------------
# Status / Resume support
# ---------------------------------------------------------------------------

def parse_skill_frontmatter(skill_path: Path) -> dict:
    """Extract YAML frontmatter metadata from a skill file."""
    text = skill_path.read_text(encoding='utf-8')
    metadata = {
        'name': skill_path.stem,
        'description': '',
        'filename': skill_path.name,
    }

    # Parse YAML frontmatter between --- delimiters
    if text.startswith('---'):
        parts = text.split('---', 2)
        if len(parts) >= 3:
            frontmatter = parts[1].strip()
            for line in frontmatter.split('\n'):
                line = line.strip()
                if line.startswith('name:'):
                    metadata['name'] = line[5:].strip().strip('"\'')
                elif line.startswith('description:'):
                    # Handle single-line description
                    desc = line[12:].strip().strip('"\'>')
                    if desc:
                        metadata['description'] = desc
                elif metadata['description'] == '' and line and not line.startswith(('---', 'name:')):
                    # Handle multi-line description (continuation lines)
                    metadata['description'] += ' ' + line.strip()

    metadata['description'] = metadata['description'].strip()
    return metadata


def install_skills(output_dir: Path, skill_paths: list[str]) -> list[dict]:
    """Copy skill files into the workspace and generate a skills manifest."""
    skills_dir = output_dir / "skills"
    skills_dir.mkdir(parents=True, exist_ok=True)

    import shutil
    installed = []

    for skill_path_str in skill_paths:
        skill_path = Path(skill_path_str)
        if not skill_path.exists():
            print(f"⚠️  Skill file not found: {skill_path}")
            continue

        # Parse metadata
        metadata = parse_skill_frontmatter(skill_path)

        # Copy to skills directory
        dest = skills_dir / skill_path.name
        shutil.copy2(skill_path, dest)

        metadata['path'] = f"skills/{skill_path.name}"
        metadata['char_count'] = len(skill_path.read_text(encoding='utf-8'))
        installed.append(metadata)

    # Write skills manifest
    if installed:
        manifest = {
            'skills': installed,
            'installed_at': datetime.now(timezone.utc).isoformat(),
        }
        manifest_path = skills_dir / "manifest.json"
        manifest_path.write_text(json.dumps(manifest, indent=2), encoding='utf-8')

    return installed


def get_review_status(output_dir: Path) -> dict:
    """Check review progress across all provision folders."""
    provisions_dir = output_dir / "provisions"
    if not provisions_dir.exists():
        return {"total": 0, "reviewed": 0, "pending": 0, "provisions": []}

    statuses = []
    for folder in sorted(provisions_dir.iterdir()):
        if not folder.is_dir():
            continue
        manifest_path = folder / "manifest.json"
        if manifest_path.exists():
            manifest = json.loads(manifest_path.read_text())
            statuses.append({
                "folder": folder.name,
                "title": manifest.get("title", "Unknown"),
                "status": manifest.get("status", "pending"),
                "reviewed_at": manifest.get("reviewed_at"),
            })

    reviewed = sum(1 for s in statuses if s["status"] == "reviewed")
    return {
        "total": len(statuses),
        "reviewed": reviewed,
        "pending": len(statuses) - reviewed,
        "provisions": statuses,
    }


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Prepare a loan agreement for agentic review with Claude Code.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Auto-detect article pattern from a text file
  python prepare_deal.py agreement.txt --posture borrower_friendly

  # Split a Word doc by heading style
  python prepare_deal.py agreement.docx --style "Heading 1" --posture balanced

  # Split by custom pattern
  python prepare_deal.py agreement.docx --pattern "ARTICLE" --posture lender_friendly

  # Check review progress
  python prepare_deal.py --status ./my_deal_review/

  # Add client-specific notes
  python prepare_deal.py agreement.txt --posture borrower_friendly \\
      --notes "Client wants aggressive push on cure periods and cash management"
        """,
    )

    parser.add_argument('input_file', nargs='?', help='Path to agreement (.docx or .txt)')
    parser.add_argument('--style', '-s', help='Heading style to split on (for .docx)')
    parser.add_argument('--pattern', '-p', help='Regex pattern to split on')
    parser.add_argument('--posture', default='borrower_friendly',
                        choices=['borrower_friendly', 'lender_friendly', 'balanced', 'compliance_only'],
                        help='Review posture (default: borrower_friendly)')
    parser.add_argument('--output-dir', '-o', help='Output directory (default: ./deal_review/)')
    parser.add_argument('--notes', '-n', default='', help='Client-specific review notes')
    parser.add_argument('--term-sheet', '-t', help='Path to term sheet or deal summary (.txt, .docx, .pdf)')
    parser.add_argument('--skill', '-k', action='append', default=[],
                        help='Path to a topical skill/reference file (.md). Repeatable: --skill file1.md --skill file2.md')
    parser.add_argument('--status', nargs='?', const='./deal_review/',
                        help='Show review progress for an existing deal directory')

    args = parser.parse_args()

    # Status check mode
    if args.status:
        status_dir = Path(args.status)
        if not status_dir.exists():
            print(f"Error: Directory not found: {status_dir}")
            return 1
        status = get_review_status(status_dir)
        print(f"\n{'='*60}")
        print(f"Deal Review Status: {status_dir}")
        print(f"{'='*60}")
        print(f"Total provisions:    {status['total']}")
        print(f"Reviewed:            {status['reviewed']}")
        print(f"Pending:             {status['pending']}")
        print(f"\nProvision Details:")
        print(f"{'-'*60}")
        for p in status['provisions']:
            icon = "✅" if p['status'] == 'reviewed' else "⏳"
            print(f"  {icon} {p['folder']}: {p['status']}")
        return 0

    if not args.input_file:
        parser.print_help()
        return 1

    input_path = Path(args.input_file)
    if not input_path.exists():
        print(f"Error: File not found: {input_path}")
        return 1

    # Determine output directory
    output_dir = Path(args.output_dir) if args.output_dir else Path('./deal_review')
    output_dir.mkdir(parents=True, exist_ok=True)

    print(f"\n{'='*60}")
    print(f"Markup — Preparing workspace")
    print(f"{'='*60}")
    print(f"Input:    {input_path}")
    print(f"Output:   {output_dir}")
    print(f"Posture:  {args.posture}")
    print()

    # --- Extract full text and split ---
    is_docx = input_path.suffix.lower() == '.docx'

    if is_docx:
        full_text = extract_full_text_from_docx(str(input_path))
    else:
        full_text = input_path.read_text(encoding='utf-8')

    # Write full agreement text
    full_agreement_path = output_dir / "full_agreement.txt"
    full_agreement_path.write_text(full_text, encoding='utf-8')
    print(f"✅ Full agreement text extracted ({len(full_text):,} chars)")

    # Compute hash for integrity tracking
    agreement_hash = hashlib.sha256(full_text.encode()).hexdigest()[:16]

    # Split into provisions
    provisions = []
    if is_docx:
        if args.style:
            provisions = split_docx_by_style(str(input_path), args.style)
        elif args.pattern:
            provisions = split_docx_by_pattern(str(input_path), args.pattern)
        else:
            # Try auto-detect on the extracted text
            detected = detect_split_pattern(full_text)
            if detected:
                print(f"📎 Auto-detected pattern: {detected}")
                provisions = split_docx_by_pattern(str(input_path), detected)
            else:
                # Fall back to trying common heading styles
                try:
                    from docx import Document
                    doc = Document(str(input_path))
                    for style in ["Heading 1", "heading 1", "HEADING 1"]:
                        provisions = split_docx_by_style(str(input_path), style)
                        if provisions:
                            print(f"📎 Split by style: {style}")
                            break
                except Exception:
                    pass
    else:
        if args.pattern:
            provisions = split_text_by_pattern(full_text, args.pattern)
        else:
            detected = detect_split_pattern(full_text)
            if detected:
                print(f"📎 Auto-detected pattern: {detected}")
                provisions = split_text_by_pattern(full_text, detected)

    if not provisions:
        print("\n⚠️  Could not detect provision boundaries.")
        print("    Try specifying --pattern or --style explicitly.")
        print("    Common patterns: 'ARTICLE', '^SECTION \\d+', '^\\d+\\.'")
        # Create a single provision with the full text as fallback
        provisions = [{
            'title': 'Full Agreement',
            'text': full_text,
            'start_line': 0,
            'section_number': '01',
        }]
        print("    Created single provision with full agreement text.\n")

    # Create provision folders
    print(f"\n📂 Creating {len(provisions)} provision folders...")
    for provision in provisions:
        folder = create_provision_folder(output_dir, provision, agreement_hash)
        word_count = len(provision['text'].split())
        cross_refs = len(detect_cross_references(provision['text']))
        print(f"   └── {folder.name} ({word_count:,} words, {cross_refs} cross-refs)")

    # Copy original file for reference
    import shutil
    original_copy = output_dir / f"original{input_path.suffix}"
    shutil.copy2(input_path, original_copy)
    print(f"✅ Original file copied to workspace")

    # --- Term sheet processing ---
    if args.term_sheet:
        term_sheet_path = Path(args.term_sheet)
        if not term_sheet_path.exists():
            print(f"⚠️  Term sheet not found: {term_sheet_path}")
        else:
            # Extract term sheet text
            ts_ext = term_sheet_path.suffix.lower()
            if ts_ext == '.docx':
                term_sheet_text = extract_full_text_from_docx(str(term_sheet_path))
            elif ts_ext == '.txt' or ts_ext == '.md':
                term_sheet_text = term_sheet_path.read_text(encoding='utf-8')
            else:
                # Copy as-is for other formats; agent can read it
                term_sheet_text = None

            # Save extracted text
            if term_sheet_text:
                ts_out = output_dir / "term_sheet.txt"
                ts_out.write_text(term_sheet_text, encoding='utf-8')
                print(f"✅ Term sheet extracted ({len(term_sheet_text):,} chars)")
            
            # Always copy the original term sheet file
            ts_copy = output_dir / f"term_sheet_original{term_sheet_path.suffix}"
            shutil.copy2(term_sheet_path, ts_copy)
            print(f"✅ Term sheet original copied")

    # --- Install topical skills ---
    if args.skill:
        installed = install_skills(output_dir, args.skill)
        if installed:
            print(f"\n📚 Installed {len(installed)} skill(s):")
            for s in installed:
                print(f"   └── {s['name']}: {s['description'][:80]}...")

    # Create review config (after term sheet and skills are processed)
    has_ts = bool(args.term_sheet and Path(args.term_sheet).exists())
    installed_skills = []
    skills_manifest = output_dir / "skills" / "manifest.json"
    if skills_manifest.exists():
        installed_skills = json.loads(skills_manifest.read_text())['skills']
    create_review_config(output_dir, args.posture, args.notes,
                        has_term_sheet=has_ts, skills=installed_skills)
    print(f"\n✅ Review config created (posture: {args.posture})")

    # --- Unpack .docx for tracked changes workflow ---
    if is_docx:
        unpacked_dir = output_dir / "unpacked"
        docx_scripts = Path("/mnt/skills/public/docx/scripts")
        if docx_scripts.exists():
            import subprocess
            unpack_script = docx_scripts / "office" / "unpack.py"
            result = subprocess.run(
                ["python3", str(unpack_script), str(input_path), str(unpacked_dir)],
                capture_output=True, text=True
            )
            if result.returncode == 0:
                print(f"✅ DOCX unpacked for tracked changes workflow")
            else:
                print(f"⚠️  DOCX unpack failed: {result.stderr[:200]}")
        else:
            print(f"ℹ️  DOCX skill not available; tracked changes workflow disabled")

    # Copy methodology as deal-level CLAUDE.md and commands into workspace
    script_dir = Path(__file__).resolve().parent.parent
    # Use the deal-review-methodology skill as the deal CLAUDE.md (contains full
    # review methodology). Fall back to root CLAUDE.md if skill not found.
    methodology_skill = script_dir / "skills" / "deal-review-methodology" / "SKILL.md"
    claude_md = script_dir / "CLAUDE.md"
    if methodology_skill.exists():
        # Strip YAML frontmatter (--- ... ---) before writing as CLAUDE.md
        content = methodology_skill.read_text(encoding="utf-8")
        if content.startswith("---"):
            end = content.find("---", 3)
            if end != -1:
                content = content[end + 3:].lstrip("\n")
        (output_dir / "CLAUDE.md").write_text(content, encoding="utf-8")
        print(f"✅ CLAUDE.md copied to workspace (from deal-review-methodology skill)")
    elif claude_md.exists():
        shutil.copy2(claude_md, output_dir / "CLAUDE.md")
        print(f"✅ CLAUDE.md copied to workspace")

    # Plugin structure: commands live in commands/ (not .claude/commands/)
    commands_src = script_dir / "commands"
    if not commands_src.exists():
        commands_src = script_dir / ".claude" / "commands"  # fallback for old structure
    if commands_src.exists():
        commands_dst = output_dir / ".claude" / "commands"
        commands_dst.mkdir(parents=True, exist_ok=True)
        for cmd_file in commands_src.glob("*.md"):
            shutil.copy2(cmd_file, commands_dst / cmd_file.name)
        print(f"✅ Slash commands copied ({len(list(commands_src.glob('*.md')))} commands)")

    # Copy scripts into workspace for assemble_deal.py access
    scripts_dst = output_dir / "scripts"
    scripts_dst.mkdir(parents=True, exist_ok=True)
    for script_file in script_dir.glob("scripts/*.py"):
        shutil.copy2(script_file, scripts_dst / script_file.name)
    print(f"✅ Scripts copied to workspace")

    # Summary
    total_words = sum(len(p['text'].split()) for p in provisions)
    print(f"\n{'='*60}")
    print(f"Workspace ready!")
    print(f"{'='*60}")
    print(f"  Provisions:  {len(provisions)}")
    print(f"  Total words:  {total_words:,}")
    print(f"  Agreement hash: {agreement_hash}")
    print(f"\nNext steps:")
    print(f"  1. cd {output_dir}")
    print(f"  2. Launch Claude Code: claude")
    print(f"  3. Run: /review-all")
    print(f"     or review a single provision: /review-provision provisions/01_...")
    print()

    return 0


if __name__ == '__main__':
    sys.exit(main())
