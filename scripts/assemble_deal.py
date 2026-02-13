#!/usr/bin/env python3
"""
assemble_deal.py — Post-processing script for Markup.

Assembles revised provisions back into a complete agreement and generates
client-facing deliverables: summary memo, changes tracker, and open issues list.

Usage:
    python scripts/assemble_deal.py ./deal_review/
    python scripts/assemble_deal.py ./deal_review/ --format docx
    python scripts/assemble_deal.py ./deal_review/ --memo-only
"""

import argparse
import json
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional


def load_provisions(deal_dir: Path) -> list[dict]:
    """Load all provisions in order with their review status and content."""
    provisions_dir = deal_dir / "provisions"
    if not provisions_dir.exists():
        return []

    provisions = []
    for folder in sorted(provisions_dir.iterdir()):
        if not folder.is_dir():
            continue

        manifest_path = folder / "manifest.json"
        if not manifest_path.exists():
            continue

        manifest = json.loads(manifest_path.read_text(encoding='utf-8'))

        provision = {
            "folder": folder.name,
            "folder_path": folder,
            "manifest": manifest,
            "original_text": None,
            "revised_text": None,
            "analysis": None,
            "changes_summary": None,
        }

        # Load available content
        for filename, key in [
            ("original.txt", "original_text"),
            ("revised.txt", "revised_text"),
            ("analysis.md", "analysis"),
            ("changes_summary.md", "changes_summary"),
        ]:
            file_path = folder / filename
            if file_path.exists():
                provision[key] = file_path.read_text(encoding='utf-8')

        provisions.append(provision)

    return provisions


def assemble_revised_agreement(provisions: list[dict]) -> str:
    """Combine all revised (or original) provisions into a single document."""
    parts = []
    for p in provisions:
        # Use revised text if available, otherwise original
        text = p["revised_text"] or p["original_text"] or ""
        if text:
            # Remove inline revision markers for the clean version
            import re
            clean_text = re.sub(r'\[REVISED:.*?\]', '', text)
            parts.append(clean_text.strip())

    return '\n\n'.join(parts)


def assemble_redline_agreement(provisions: list[dict]) -> str:
    """Combine all provisions keeping revision markers visible."""
    parts = []
    for p in provisions:
        text = p["revised_text"] or p["original_text"] or ""
        if text:
            parts.append(text.strip())

    return '\n\n'.join(parts)


def generate_review_memo(
    deal_dir: Path,
    provisions: list[dict],
    config: dict,
) -> str:
    """Generate a client-facing review summary memo."""
    posture = config.get("review_posture", "unknown")
    now = datetime.now(timezone.utc).strftime("%B %d, %Y")

    total = len(provisions)
    reviewed = sum(1 for p in provisions if p["manifest"].get("status") == "reviewed")
    pending = total - reviewed

    # Collect all cross-reference flags
    all_flags = []
    all_open_issues = []
    for p in provisions:
        m = p["manifest"]
        for flag in m.get("cross_ref_flags", []):
            all_flags.append(f"  - [{m['title']}] {flag}")
        for issue in m.get("open_issues", []):
            all_open_issues.append(f"  - [{m['title']}] {issue}")

    memo = f"""# DEAL REVIEW MEMORANDUM
## Confidential — Attorney Work Product

**Date:** {now}
**Review Posture:** {posture.replace('_', ' ').title()}
**Provisions Reviewed:** {reviewed} of {total}
{"**Status: INCOMPLETE — " + str(pending) + " provisions pending**" if pending > 0 else "**Status: COMPLETE**"}

---

## Executive Summary

This memorandum summarizes the results of an AI-assisted review of the loan agreement.
Each article/section was analyzed individually with access to the full agreement context.
All revisions reflect a {posture.replace('_', ' ')} posture.

**Important:** This AI-assisted review is a tool to support attorney analysis, not a
substitute for it. All revisions and recommendations should be reviewed by counsel
before delivery to any party.

---

## Provision-by-Provision Summary

"""

    for p in provisions:
        m = p["manifest"]
        status_icon = "✅" if m.get("status") == "reviewed" else "⏳"
        memo += f"### {status_icon} {m['title']}\n\n"

        if p["changes_summary"]:
            # Extract just the key points from the changes summary
            memo += p["changes_summary"] + "\n\n"
        elif m.get("status") == "reviewed":
            memo += "*Reviewed — see individual provision analysis for details.*\n\n"
        else:
            memo += "*Pending review.*\n\n"

        memo += "---\n\n"

    # Cross-reference issues
    if all_flags:
        memo += "## Cross-Reference Issues\n\n"
        memo += "The following cross-reference conflicts were identified during review:\n\n"
        memo += '\n'.join(all_flags) + "\n\n---\n\n"

    # Open issues
    if all_open_issues:
        memo += "## Open Issues Requiring Client Input\n\n"
        memo += "The following items require business-point decisions:\n\n"
        memo += '\n'.join(all_open_issues) + "\n\n---\n\n"

    memo += """## Disclaimer

This review was performed using an AI-assisted workflow. While the AI agent had access
to the full agreement for context and followed structured legal analysis methodology,
all output should be treated as a first draft requiring attorney review. Defined terms,
cross-references, and legal conclusions should be independently verified.
"""

    return memo


def generate_changes_tracker(provisions: list[dict]) -> str:
    """Generate a consolidated changes tracking document."""
    tracker = "# CHANGES TRACKER\n\n"
    tracker += "| # | Provision | Status | Changes | Open Issues |\n"
    tracker += "|---|-----------|--------|---------|-------------|\n"

    for p in provisions:
        m = p["manifest"]
        status = m.get("status", "pending")
        changes = "See analysis" if p["changes_summary"] else "—"
        issues = str(len(m.get("open_issues", []))) + " issues" if m.get("open_issues") else "None"
        tracker += f"| {m['section_number']} | {m['title'][:40]} | {status} | {changes} | {issues} |\n"

    return tracker


def write_docx(text: str, output_path: Path, title: str = "Document"):
    """Write text to a .docx file (requires python-docx)."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except ImportError:
        print("Warning: python-docx not available. Saving as .txt instead.")
        output_path = output_path.with_suffix('.txt')
        output_path.write_text(text, encoding='utf-8')
        return output_path

    doc = Document()

    # Simple conversion: split by lines and headings
    for line in text.split('\n'):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('---'):
            doc.add_paragraph('_' * 50)
        else:
            doc.add_paragraph(stripped)

    doc.save(str(output_path))
    return output_path


def main():
    parser = argparse.ArgumentParser(
        description="Assemble reviewed provisions into deliverables.",
    )
    parser.add_argument('deal_dir', help='Path to the deal review directory')
    parser.add_argument('--format', choices=['txt', 'docx'], default='txt',
                        help='Output format (default: txt)')
    parser.add_argument('--memo-only', action='store_true',
                        help='Generate only the review memo')
    parser.add_argument('--output-dir', '-o',
                        help='Output directory (default: <deal_dir>/deliverables/)')

    args = parser.parse_args()

    deal_dir = Path(args.deal_dir)
    if not deal_dir.exists():
        print(f"Error: Directory not found: {deal_dir}")
        return 1

    output_dir = Path(args.output_dir) if args.output_dir else deal_dir / "deliverables"
    output_dir.mkdir(parents=True, exist_ok=True)

    # Load config
    config_path = deal_dir / "review_config.json"
    config = {}
    if config_path.exists():
        config = json.loads(config_path.read_text(encoding='utf-8'))

    # Load provisions
    provisions = load_provisions(deal_dir)
    if not provisions:
        print("Error: No provisions found in deal directory.")
        return 1

    reviewed = sum(1 for p in provisions if p["manifest"].get("status") == "reviewed")
    print(f"\n{'='*60}")
    print(f"Assembling Deliverables")
    print(f"{'='*60}")
    print(f"Provisions: {len(provisions)} total, {reviewed} reviewed")
    print()

    ext = '.docx' if args.format == 'docx' else '.txt'

    # 1. Review Memo (always generated)
    memo = generate_review_memo(deal_dir, provisions, config)
    if args.format == 'docx':
        memo_path = write_docx(memo, output_dir / "review_memo.docx", "Review Memorandum")
    else:
        memo_path = output_dir / "review_memo.md"
        memo_path.write_text(memo, encoding='utf-8')
    print(f"✅ Review memo: {memo_path}")

    if args.memo_only:
        return 0

    # 2. Revised Agreement (clean)
    revised = assemble_revised_agreement(provisions)
    if args.format == 'docx':
        revised_path = write_docx(revised, output_dir / "revised_agreement.docx")
    else:
        revised_path = output_dir / "revised_agreement.txt"
        revised_path.write_text(revised, encoding='utf-8')
    print(f"✅ Revised agreement (clean): {revised_path}")

    # 3. Redline Agreement (with markers)
    redline = assemble_redline_agreement(provisions)
    redline_path = output_dir / "redline_agreement.txt"
    redline_path.write_text(redline, encoding='utf-8')
    print(f"✅ Redline agreement: {redline_path}")

    # 4. Changes Tracker
    tracker = generate_changes_tracker(provisions)
    tracker_path = output_dir / "changes_tracker.md"
    tracker_path.write_text(tracker, encoding='utf-8')
    print(f"✅ Changes tracker: {tracker_path}")

    # 5. Copy reconciliation report if it exists
    recon_path = deal_dir / "reconciliation_report.md"
    if recon_path.exists():
        import shutil
        dest = output_dir / "reconciliation_report.md"
        shutil.copy2(recon_path, dest)
        print(f"✅ Reconciliation report: {dest}")

    # 6. Copy term sheet compliance report if it exists
    ts_report = deal_dir / "term_sheet_compliance_report.md"
    if ts_report.exists():
        dest = output_dir / "term_sheet_compliance_report.md"
        shutil.copy2(ts_report, dest)
        print(f"✅ Term sheet compliance report: {dest}")

    # 7. Copy redline Word document if it exists
    redline_docx = deal_dir / "redline_agreement.docx"
    if redline_docx.exists():
        dest = output_dir / "redline_agreement.docx"
        shutil.copy2(redline_docx, dest)
        print(f"✅ Redline Word document: {dest}")
    elif (deal_dir / "unpacked").exists():
        print(f"ℹ️  Unpacked .docx found but no redline generated yet.")
        print(f"    Run /apply-redlines to produce a Word document with tracked changes.")

    print(f"\n{'='*60}")
    print(f"All deliverables saved to: {output_dir}")
    print(f"{'='*60}\n")

    return 0


if __name__ == '__main__':
    sys.exit(main())
